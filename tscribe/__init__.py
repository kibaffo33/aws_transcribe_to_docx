""" Transform AWS Transcribe json files to docx, csv, sqlite and vtt. """

from docx import Document
from docx.shared import Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, datetime
import matplotlib.pyplot as plt
import statistics
from pathlib import Path
from time import perf_counter
import pandas
import sqlite3
import webvtt
import logging


def convert_time_stamp(timestamp: str) -> str:
    """ Function to help convert timestamps from s to H:M:S """
    delta = datetime.timedelta(seconds=float(timestamp))
    seconds = delta - datetime.timedelta(microseconds=delta.microseconds)
    return str(seconds)


def load_json_as_dict(filepath: str) -> dict:
    """Load in JSON file and return as dict"""
    logging.info("Loading json")

    json_filepath = Path(filepath)
    assert json_filepath.is_file(), "JSON file does not exist"

    data = json.load(open(json_filepath.absolute(), "r", encoding="utf-8"))
    assert "jobName" in data
    assert "results" in data
    assert "status" in data

    assert data["status"] == "COMPLETED", "JSON file not shown as completed."

    logging.debug("json checks psased")
    return data


def calculate_confidence_statistics(data: dict) -> dict:
    """Confidence Statistics"""
    logging.info("Gathering confidence statistics")

    # Stats dictionary
    stats = {
        "timestamps": [],
        "accuracy": [],
        "9.8": 0,
        "9": 0,
        "8": 0,
        "7": 0,
        "6": 0,
        "5": 0,
        "4": 0,
        "3": 0,
        "2": 0,
        "1": 0,
        "0": 0,
        "total": len(data["results"]["items"]),
    }

    # Confidence count
    for item in data["results"]["items"]:
        if item["type"] == "pronunciation":

            stats["timestamps"].append(float(item["start_time"]))

            confidence_decimal = float(item["alternatives"][0]["confidence"])
            confidence_integer = int(confidence_decimal * 100)

            stats["accuracy"].append(confidence_integer)

            if confidence_decimal >= 0.98:
                stats["9.8"] += 1
            else:
                rough_confidence = str(int(confidence_decimal * 10))
                stats[rough_confidence] += 1

    return stats


def make_graph_png(stats: dict, directory: str) -> str:
    """Make scatter graph from confidence statistics"""
    logging.info("Making graph")

    # Confidence of each word as scatter graph
    plt.scatter(stats["timestamps"], stats["accuracy"])

    # Mean average as line across graph
    plt.plot(
        [stats["timestamps"][0], stats["timestamps"][-1]],
        [statistics.mean(stats["accuracy"]), statistics.mean(stats["accuracy"])],
        "r",
    )

    # Formatting
    plt.xlabel("Time (seconds)")
    plt.ylabel("Accuracy (percent)")
    plt.yticks(range(0, 101, 10))
    plt.title("Accuracy during transcript")
    plt.legend(["Accuracy average (mean)", "Individual words"], loc="lower center")

    # Target filename, including directory for explicit path
    filename = Path(directory) / Path("chart.png")
    plt.savefig(str(filename))
    logging.info("Graph saved to %s", filename)
    plt.clf()

    return str(filename)


def decode_transcript_to_dataframe(data: str):
    """Decode the transcript into a pandas dataframe"""
    logging.info("Decoding transcript")

    decoded_data = {"start_time": [], "end_time": [], "speaker": [], "comment": []}

    # If speaker identification
    if "speaker_labels" in data["results"].keys():
        logging.debug("Transcipt has speaker_labels")

        # A segment is a blob of pronounciation and punctuation by an individual speaker
        for segment in data["results"]["speaker_labels"]["segments"]:

            # If there is content in the segment, add a row, write the time and speaker
            if len(segment["items"]) > 0:
                decoded_data["start_time"].append(
                    convert_time_stamp(segment["start_time"])
                )
                decoded_data["end_time"].append(convert_time_stamp(segment["end_time"]))
                decoded_data["speaker"].append(segment["speaker_label"])
                decoded_data["comment"].append("")

                # For each word in the segment...
                for word in segment["items"]:

                    # Get the word with the highest confidence
                    pronunciations = list(
                        filter(
                            lambda x: x["type"] == "pronunciation",
                            data["results"]["items"],
                        )
                    )
                    word_result = list(
                        filter(
                            lambda x: x["start_time"] == word["start_time"]
                            and x["end_time"] == word["end_time"],
                            pronunciations,
                        )
                    )
                    result = sorted(
                        word_result[-1]["alternatives"], key=lambda x: x["confidence"]
                    )[-1]

                    # Write the word
                    decoded_data["comment"][-1] += " " + result["content"]

                    # If the next item is punctuation, write it
                    try:
                        word_result_index = data["results"]["items"].index(
                            word_result[0]
                        )
                        next_item = data["results"]["items"][word_result_index + 1]
                        if next_item["type"] == "punctuation":
                            decoded_data["comment"][-1] += next_item["alternatives"][0][
                                "content"
                            ]
                    except IndexError:
                        pass

    # If channel identification
    elif "channel_labels" in data["results"].keys():
        logging.debug("Transcipt has channel_labels")

        # For each word in the results
        for word in data["results"]["items"]:

            # Punctuation items do not include a start_time
            if "start_time" not in word.keys():
                continue

            # Identify the channel
            channel = list(
                filter(
                    lambda x: word in x["items"],
                    data["results"]["channel_labels"]["channels"],
                )
            )[0]["channel_label"]

            # If still on the same channel, add the current word to the line
            if (
                channel in decoded_data["speaker"]
                and decoded_data["speaker"][-1] == channel
            ):
                current_word = sorted(
                    word["alternatives"], key=lambda x: x["confidence"]
                )[-1]
                decoded_data["comment"][-1] += " " + current_word["content"]

            # Else start a new line
            else:
                decoded_data["start_time"].append(
                    convert_time_stamp(word["start_time"])
                )
                decoded_data["end_time"].append(convert_time_stamp(word["end_time"]))
                decoded_data["speaker"].append(channel)
                current_word = sorted(
                    word["alternatives"], key=lambda x: x["confidence"]
                )[-1]
                decoded_data["comment"].append(current_word["content"])

            # If the next item is punctuation, write it
            try:
                word_result_index = data["results"]["items"].index(word)
                next_item = data["results"]["items"][word_result_index + 1]
                if next_item["type"] == "punctuation":
                    decoded_data["comment"][-1] += next_item["alternatives"][0][
                        "content"
                    ]
            except IndexError:
                pass

    # Neither speaker nor channel identification
    else:
        logging.debug("No speaker_labels or channel_labels")

        decoded_data["start_time"] = convert_time_stamp(
            list(
                filter(lambda x: x["type"] == "pronunciation", data["results"]["items"])
            )[0]["start_time"]
        )
        decoded_data["end_time"] = convert_time_stamp(
            list(
                filter(lambda x: x["type"] == "pronunciation", data["results"]["items"])
            )[-1]["end_time"]
        )
        decoded_data["speaker"].append("")
        decoded_data["comment"].append("")

        # Add words
        for word in data["results"]["items"]:

            # Get the word with the highest confidence
            result = sorted(word["alternatives"], key=lambda x: x["confidence"])[-1]

            # Write the word
            decoded_data["comment"][-1] += " " + result["content"]

            # If the next item is punctuation, write it
            try:
                word_result_index = data["results"]["items"].index(word)
                next_item = data["results"]["items"][word_result_index + 1]
                if next_item["type"] == "punctuation":
                    decoded_data["comment"][-1] += next_item["alternatives"][0][
                        "content"
                    ]
            except IndexError:
                pass

    # Produce pandas dataframe
    dataframe = pandas.DataFrame(
        decoded_data, columns=["start_time", "end_time", "speaker", "comment"]
    )

    # Clean leading whitespace
    dataframe["comment"] = dataframe["comment"].str.lstrip()

    return dataframe


def write_docx(data, filename, **kwargs):
    """ Write a transcript from the .json transcription file. """
    logging.info("Writing docx")

    output_filename = Path(filename)

    # Initiate Document
    document = Document()
    # A4 Size
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)
    # Font
    font = document.styles["Normal"].font
    font.name = "Calibri"

    # Document title and intro
    title = f"Transcription of {data['jobName']}"
    document.add_heading(title, level=1)
    # Set thresholds for formatting later
    threshold_for_grey = 0.98
    grey = RGBColor(204, 204, 204)
    threshold_for_orange = 0.75
    orange = RGBColor(255, 114, 0)
    threshold_for_brown = 0.50
    brown = RGBColor(156, 109, 7)
    threshold_for_purple = 0.25
    purple = RGBColor(106, 90, 205)
    threshold_for_red = 0.00
    red = RGBColor(255, 0, 0)
    # Intro
    document.add_paragraph(
        "Transcription using AWS Transcribe automatic speech recognition and"
        " the 'tscribe' python package."
    )
    document.add_paragraph(
        datetime.datetime.now().strftime("Document produced on %A %d %B %Y at %X.")
    )
    document.add_paragraph()  # Spacing
    document.add_paragraph(
        f"Grey text has less than {int(threshold_for_grey * 100)}% confidence."
    )
    document.add_paragraph(
        f"Orange text has less than {int(threshold_for_orange * 100)}% confidence."
    )
    document.add_paragraph(
        f"Brown text has less than {int(threshold_for_brown * 100)}% confidence."
    )
    document.add_paragraph(
        f"Purple text has less than {int(threshold_for_purple * 100)}% confidence."
    )
    document.add_paragraph(
        f"Red text has less than {int(threshold_for_red * 100)}% confidence."
    )

    # Get stats
    stats = calculate_confidence_statistics(data)

    # Display confidence count table
    table = document.add_table(rows=1, cols=3)
    table.style = document.styles["Light List Accent 1"]
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Confidence"
    hdr_cells[1].text = "Count"
    hdr_cells[2].text = "Percentage"
    row_cells = table.add_row().cells
    row_cells[0].text = str("98% - 100%")
    row_cells[1].text = str(stats["9.8"])
    row_cells[2].text = str(round(stats["9.8"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("90% - 97%")
    row_cells[1].text = str(stats["9"])
    row_cells[2].text = str(round(stats["9"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("80% - 89%")
    row_cells[1].text = str(stats["8"])
    row_cells[2].text = str(round(stats["8"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("70% - 79%")
    row_cells[1].text = str(stats["7"])
    row_cells[2].text = str(round(stats["7"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("60% - 69%")
    row_cells[1].text = str(stats["6"])
    row_cells[2].text = str(round(stats["6"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("50% - 59%")
    row_cells[1].text = str(stats["5"])
    row_cells[2].text = str(round(stats["5"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("40% - 49%")
    row_cells[1].text = str(stats["4"])
    row_cells[2].text = str(round(stats["4"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("30% - 39%")
    row_cells[1].text = str(stats["3"])
    row_cells[2].text = str(round(stats["3"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("20% - 29%")
    row_cells[1].text = str(stats["2"])
    row_cells[2].text = str(round(stats["2"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("10% - 19%")
    row_cells[1].text = str(stats["1"])
    row_cells[2].text = str(round(stats["1"] / stats["total"] * 100, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = str("0% - 9%")
    row_cells[1].text = str(stats["0"])
    row_cells[2].text = str(round(stats["0"] / stats["total"] * 100, 2)) + "%"
    # Add paragraph for spacing
    document.add_paragraph()

    graph = make_graph_png(stats, str(output_filename.parent))
    document.add_picture(graph, width=Cm(14.64))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    # Process and display transcript by speaker segments
    table = document.add_table(rows=1, cols=3)
    table.style = document.styles["Light List Accent 1"]
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Time"
    hdr_cells[1].text = "Speaker"
    hdr_cells[2].text = "Content"

    # If speaker identification
    if "speaker_labels" in data["results"].keys():
        logging.debug("Transcript has speaker_labels")

        # A segment is a blob of pronounciation and punctuation by an individual speaker
        for segment in data["results"]["speaker_labels"]["segments"]:

            # If there is content in the segment, add a row, write the time and speaker
            if len(segment["items"]) > 0:
                row_cells = table.add_row().cells
                row_cells[0].text = convert_time_stamp(segment["start_time"])
                row_cells[1].text = str(segment["speaker_label"])

                # For each word in the segment...
                for word in segment["items"]:

                    # Get the word with the highest confidence
                    pronunciations = list(
                        filter(
                            lambda x: x["type"] == "pronunciation",
                            data["results"]["items"],
                        )
                    )
                    word_result = list(
                        filter(
                            lambda x: x["start_time"] == word["start_time"]
                            and x["end_time"] == word["end_time"],
                            pronunciations,
                        )
                    )
                    result = sorted(
                        word_result[-1]["alternatives"], key=lambda x: x["confidence"]
                    )[-1]

                    # Write the word
                    run = row_cells[2].paragraphs[0].add_run(" " + result["content"])
                    if float(result["confidence"]) < threshold_for_red:
                        font = run.font
                        font.color.rgb = red
                    elif float(result["confidence"]) < threshold_for_purple:
                        font = run.font
                        font.color.rgb = purple
                    elif float(result["confidence"]) < threshold_for_brown:
                        font = run.font
                        font.color.rgb = brown
                    elif float(result["confidence"]) < threshold_for_orange:
                        font = run.font
                        font.color.rgb = orange
                    elif float(result["confidence"]) < threshold_for_grey:
                        font = run.font
                        font.color.rgb = grey

                    # If the next item is punctuation, write it
                    try:
                        word_result_index = data["results"]["items"].index(
                            word_result[0]
                        )
                        next_item = data["results"]["items"][word_result_index + 1]
                        if next_item["type"] == "punctuation":
                            run = (
                                row_cells[2]
                                .paragraphs[0]
                                .add_run(next_item["alternatives"][0]["content"])
                            )
                    except IndexError:
                        pass

    # If channel identification
    elif "channel_labels" in data["results"].keys():
        logging.debug("Transcript has channel_labels")

        for word in data["results"]["items"]:

            # Punctuation items do not include a start_time
            if "start_time" not in word.keys():
                continue

            # Identify the channel
            channel = list(
                filter(
                    lambda x: word in x["items"],
                    data["results"]["channel_labels"]["channels"],
                )
            )[0]["channel_label"]

            # If still on the same channel, add the current word to the line
            if table.cell(-1, 1).text == channel:
                current_word = sorted(
                    word["alternatives"], key=lambda x: x["confidence"]
                )[-1]

                run = (
                    table.cell(-1, 2)
                    .paragraphs[0]
                    .add_run(" " + current_word["content"])
                )
                
                if float(result["confidence"]) < threshold_for_red:
                    font = run.font
                    font.color.rgb = red
                elif float(result["confidence"]) < threshold_for_purple:
                    font = run.font
                    font.color.rgb = purple
                elif float(result["confidence"]) < threshold_for_brown:
                    font = run.font
                    font.color.rgb = brown
                elif float(result["confidence"]) < threshold_for_orange:
                    font = run.font
                    font.color.rgb = orange
                elif float(result["confidence"]) < threshold_for_grey:
                    font = run.font
                    font.color.rgb = grey

            # Else start a new line
            else:
                current_word = sorted(
                    word["alternatives"], key=lambda x: x["confidence"]
                )[-1]

                row_cells = table.add_row().cells
                row_cells[0].text = convert_time_stamp(word["start_time"])
                row_cells[1].text = channel

                run = row_cells[2].paragraphs[0].add_run(" " + current_word["content"])
                
                if float(result["confidence"]) < threshold_for_red:
                        font = run.font
                        font.color.rgb = red
                elif float(result["confidence"]) < threshold_for_purple:
                    font = run.font
                    font.color.rgb = purple
                elif float(result["confidence"]) < threshold_for_brown:
                    font = run.font
                    font.color.rgb = brown
                elif float(result["confidence"]) < threshold_for_orange:
                    font = run.font
                    font.color.rgb = orange
                elif float(result["confidence"]) < threshold_for_grey:
                    font = run.font
                    font.color.rgb = grey
            # If the next item is punctuation, write it
            try:
                word_result_index = data["results"]["items"].index(word)
                next_item = data["results"]["items"][word_result_index + 1]
                if next_item["type"] == "punctuation":
                    run = (
                        row_cells[2]
                        .paragraphs[0]
                        .add_run(next_item["alternatives"][0]["content"])
                    )
            except IndexError:
                pass

    # Else no speaker identification
    else:
        logging.debug("No speaker_labels or channel_labels")

        # Start the first row
        row_cells = table.add_row().cells

        # Add words
        for word in data["results"]["items"]:

            # Get the word with the highest confidence
            result = sorted(word["alternatives"], key=lambda x: x["confidence"])[-1]

            # Write the word
            run = row_cells[2].paragraphs[0].add_run(" " + result["content"])
            if float(result["confidence"]) < threshold_for_red:
                font = run.font
                font.color.rgb = red
            elif float(result["confidence"]) < threshold_for_purple:
                font = run.font
                font.color.rgb = purple
            elif float(result["confidence"]) < threshold_for_brown:
                font = run.font
                font.color.rgb = brown
            elif float(result["confidence"]) < threshold_for_orange:
                font = run.font
                font.color.rgb = orange
            elif float(result["confidence"]) < threshold_for_grey:
                font = run.font
                font.color.rgb = grey

            # If the next item is punctuation, write it
            try:
                word_result_index = data["results"]["items"].index(word)
                next_item = data["results"]["items"][word_result_index + 1]
                if next_item["type"] == "punctuation":
                    run = (
                        row_cells[2]
                        .paragraphs[0]
                        .add_run(next_item["alternatives"][0]["content"])
                    )
            except IndexError:
                pass

    # Formatting transcript table widthds
    widths = (Inches(0.6), Inches(1), Inches(4.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Save
    document.save(filename)
    logging.info("Docx saved to %s", filename)


def write_vtt(dataframe, filename):
    """Output to VTT format"""
    logging.info("Writing VTT")

    # Initialize vtt
    vtt = webvtt.WebVTT()

    # Iterate through dataframe
    for _, row in dataframe.iterrows():

        # If the segment has 80 or less characters
        if len(row["comment"]) <= 80:

            caption = webvtt.Caption(
                start=row["start_time"] + ".000",
                end=row["end_time"] + ".000",
                text=row["comment"],
            )

        # If the segment has more than 80 characters, use lines
        else:

            lines = []
            text = row["comment"]

            while len(text) > 80:
                text = text.lstrip()
                last_space = text[:80].rindex(" ")
                lines.append(text[:last_space])
                text = text[last_space:]

            caption = webvtt.Caption(
                row["start_time"] + ".000", row["end_time"] + ".000", lines
            )

        if row["speaker"]:
            caption.identifier = row["speaker"]

        vtt.captions.append(caption)

    vtt.save(filename)
    logging.info("VTT saved to %s", filename)


def write(transcript_filepath, **kwargs):
    """Main function, write transcript file from json"""

    # Performance timer start
    start = perf_counter()
    logging.info("=" * 32)
    logging.debug("Started at %s", start)
    logging.info("Source file: %s", transcript_filepath)
    logging.debug("kwargs = %s", str(kwargs))

    # Load json file as dict
    data = load_json_as_dict(transcript_filepath)

    # Decode transcript
    dataframe = decode_transcript_to_dataframe(data)

    # Output
    output_format = kwargs.get("format", "docx")

    # Deprecated tmp_dir by improving save_as
    if kwargs.get("tmp_dir"):
        logging.warning("tmp_dir in kwargs")
        raise Exception("tmp_dir has been deprecated, use save_as instead")

    # Output to docx (default behaviour)
    if output_format == "docx":
        output_filepath = kwargs.get(
            "save_as", Path(transcript_filepath).with_suffix(".docx")
        )
        write_docx(data, output_filepath)

    # Output to CSV
    elif output_format == "csv":
        output_filepath = kwargs.get(
            "save_as", Path(transcript_filepath).with_suffix(".csv")
        )
        dataframe.to_csv(output_filepath)

    # Output to sqlite
    elif output_format == "sqlite":
        output_filepath = kwargs.get(
            "save_as", Path(transcript_filepath).with_suffix(".db")
        )
        conn = sqlite3.connect(str(output_filepath))
        dataframe.to_sql("transcript", conn)
        conn.close()

    # Output to VTT
    elif output_format == "vtt":
        output_filepath = kwargs.get(
            "save_as", Path(transcript_filepath).with_suffix(".vtt")
        )
        write_vtt(dataframe, output_filepath)

    else:
        raise Exception("Output format should be 'docx', 'csv', 'sqlite' or 'vtt'")

    # Performance timer finish
    finish = perf_counter()
    logging.debug("Finished at %s", finish)
    duration = round(finish - start, 2)

    print(f"{output_filepath} written in {duration} seconds.")
    logging.info("%s written in %s seconds.", output_filepath, duration)

import pytest
import tscribe
import os
import pandas
from uuid import uuid4
from pathlib import Path
import sqlite3
from docx import Document


sample_files = [
    "sample_single.json", 
    "sample_multiple.json",]
"""
    "B1.json",
    "B2.json",
    "B3.json",
    "B4.json",
    "B5.json",
    "S1.json",
    "S2.json",
    "S3.json",
    "S4.json",
    "S5.json",
    "V1.json",
    "V2.json",
    "V3.json",
    "V4.json",
    "V5.json",
    "V5_2speakers_nodict.json"
]
"""


@pytest.mark.parametrize("time_stamp,expected", [("1.0", "0:00:01"), ("2.5", "0:00:02"), ("60.0", "0:01:00"), ("3600", "1:00:00")])
def test_convert_time_stamp(time_stamp, expected):
    """
    Test timetsamp conversion utility function
    
    GIVEN a float of seconds as data type str
    WHEN calling convert_time_stamp(...)
    THEN convert the float of seconds to a H:MM:SS format
    """

    # GIVEN a float of seconds as data type str
    time_stamp = time_stamp

    # WHEN calling convert_time_stamp(...)
    result = tscribe.convert_time_stamp(time_stamp)

    # THEN convert the float of seconds to a H:MM:SS format
    assert result == expected, f"Result of {time_stamp} should be {expected}"

    result = result.split(":")
    seconds = int(result[2])
    minutes = int(result[1]) * 60
    hours = int(result[0]) * 60 * 60
    total_seconds = seconds + minutes + hours
    assert int(float(time_stamp)) == total_seconds, f"Reverse calculation of {time_stamp} shoud be {total_seconds}"


@pytest.mark.parametrize("input_file", sample_files)
def test_load_json(input_file):
    """
    Test json to dict function
    
    GIVEN a sample json file
    WHEN calling tscribe.load_json(...)
    THEN return a dict
    """

    # GIVEN a sample json file
    # provided through parametrize

    # WHEN calling tscribe.load_json(...)
    data = tscribe.load_json(input_file)

    # THEN return a dict
    assert isinstance(data, dict), "Data should by of dict type"


@pytest.mark.parametrize("input_file", sample_files)
def test_confidence_stats(input_file):
    """
    Test confidence stats data modeling

    GIVEN a data dict
    WHEN calling confidence_stats(...)
    THEN return the data model with the right components
    """

    # GIVEN a data dict
    # input_file = "sample_multiple.json"
    data = tscribe.load_json(input_file)

    # WHEN calling confidence_stats(...)
    stats = tscribe.confidence_stats(data)

    # THEN return the data model with the right components
    assert isinstance(stats, dict), "Stats should be of dict type"
    assert "timestamps" in stats, "Data model should include timestamps"
    assert "9.8" in stats, "Data model should include 9.8"
    assert "9" in stats, "Data model should include 9"
    assert "8" in stats, "Data model should include 8"
    assert "7" in stats, "Data model should include 7"
    assert "6" in stats, "Data model should include 6"
    assert "5" in stats, "Data model should include 5"
    assert "4" in stats, "Data model should include 4"
    assert "3" in stats, "Data model should include 3"
    assert "2" in stats, "Data model should include 2"
    assert "1" in stats, "Data model should include 1"
    assert "0" in stats, "Data model should include 0"


@pytest.mark.parametrize("input_file", sample_files)
def test_make_graph(input_file):
    """
    Test function for creating graphs from confidence stats
    
    GIVEN confidence stats from an input file
    WHEN calling make_graph(...)
    THEN produce chart.png
    """
    
    filepath = Path("chart.png")

    # Ensure blank slate
    if filepath.is_file():
        os.remove(filepath)

    # GIVEN confidence stats from an input file
    data = tscribe.load_json(input_file)
    stats = tscribe.confidence_stats(data)

    # WHEN calling make_graph(...)
    tscribe.make_graph(stats, "./")

    # THEN produce chart.png
    assert filepath.is_file(), "chart.png should be created"

    os.remove(filepath)


@pytest.mark.parametrize("input_file", sample_files)
def test_decode_transcript(input_file):
    """
    Test transcript decoding function

    GIVEN a data dict
    WHEN calling decode_transcript(...)
    THEN 
    """

    # GIVEN a data dict
    data = tscribe.load_json(input_file)

    # WHEN calling decode_transcript(...)
    df = tscribe.decode_transcript(data)

    # THEN
    assert isinstance(df, pandas.DataFrame), "decode_transcript should return a Pandas Data Frame"

    rows, cols = df.shape
    
    assert cols == 3, "Dataframe should have three columns"

    if input_file == "sample_single.json":
        # TODO
        pass
    
    if input_file == "sample_multiple.json":
        assert rows == len(data['results']['speaker_labels']['segments']), "Rows should match number of segments"


@pytest.mark.parametrize("input_file", sample_files)
def test_write_to_docx(input_file):
    """
    Test production of docx output
    
    GIVEN an input file
    WHEN writing to docx
    THEN check output exists and contains content
    """

    # GIVEN an input file
    # WHEN writing to docx
    output_filename = Path(f"{uuid4().hex}.docx")
    tscribe.write(input_file, save_as=output_filename, format="docx")

    # THEN check output exists and contains content
    assert output_filename.is_file(), "Output file should exist"

    document = Document(output_filename)
    
    assert len(document.tables) == 2, "Document should contain two tables, stats and transcript"

    t_conf = document.tables[0].cell(0,0).text
    t_count = document.tables[0].cell(0,1).text
    t_perc = document.tables[0].cell(0,2).text
    assert (t_conf, t_count, t_perc) == ("Confidence", "Count", "Percentage"), "First table should be stats headers"
    assert len(document.tables[0].rows) == 12, "Stats table should hold 12 rows"

    t_time = document.tables[1].cell(0,0).text
    t_speaker = document.tables[1].cell(0,1).text
    t_content = document.tables[1].cell(0,2).text
    assert (t_time, t_speaker, t_content) == ("Time", "Speaker", "Content"), "Second table should be transcript headers"
    data = tscribe.load_json(input_file)
    df = tscribe.decode_transcript(data)
    assert len(document.tables[1].rows) == len(df) + 1, "Second table should be length of dataframe + headers"

    assert "chart.png" in document.paragraphs[6]._p.xml, "Chart should be in paragraph six"

    # Teardown
    os.remove(output_filename)


@pytest.mark.parametrize("input_file", sample_files)
def test_write_to_csv(input_file):
    """
    Test production of csv output
    
    GIVEN an input file
    WHEN writing to csv
    THEN check output exists and contains content
    """

    # GIVEN an input file
    # WHEN writing to csv
    output_filename = Path(f"{uuid4().hex}.csv")
    tscribe.write(input_file, save_as=output_filename, format="csv")

    # THEN check output exists and contains content
    assert output_filename.is_file(), "Output file should exist"

    with open(output_filename, 'r') as file:
        lines = file.readlines()

    data = tscribe.load_json(input_file)
    df = tscribe.decode_transcript(data)

    assert len(lines) == len(df) + 1, "CSV should be length of dataframe + headers"

    # Teardown
    os.remove(output_filename)


@pytest.mark.parametrize("input_file", sample_files)
def test_write_to_sqlite(input_file):
    """
    Test production of sqlite output
    
    GIVEN an input file
    WHEN writing to sqlite
    THEN check output exists and contains content
    """

    # GIVEN an input file
    # WHEN writing to sqlite
    output_filename = Path(f"{uuid4().hex}.db")
    tscribe.write(input_file, save_as=output_filename, format="sqlite")

    # THEN check output exists and contains content
    assert output_filename.is_file(), "Output file should exist"

    conn = sqlite3.connect(output_filename)
    c = conn.cursor()
    c.execute("SELECT * FROM transcript")
    query = c.fetchall()

    data = tscribe.load_json(input_file)
    df = tscribe.decode_transcript(data)

    assert len(query) == len(df), "Database table should be length of dataframe"

    # Teardown
    os.remove(output_filename)


@pytest.mark.parametrize("input_file", sample_files)
def test_write_to_default(input_file):
    """
    Test production of default output
    
    GIVEN an input file
    WHEN not specifying output
    THEN check output is the default format
    """

    # GIVEN an input file
    # WHEN not specifying output
    tscribe.write(input_file)
    expected_filename = input_file.replace(".json", ".docx")
    output_filename = Path(expected_filename)

    # THEN check output exists and contains content
    assert output_filename.is_file(), "Output file should exist"

    # Teardown
    os.remove(output_filename)


@pytest.mark.parametrize("input_file", sample_files)
@pytest.mark.parametrize("output_format", ["docx", "csv", "sqlite"])
@pytest.mark.parametrize("location", [".", "output"])
def test_save_as(input_file, output_format, location):
    """
    Test saving of supported formats to locations
    
    GIVEN locations of current or specific folder
    WHEN writing transcript in any supported format
    THEN check output exists
    """

    if not Path("output").is_dir():
        os.mkdir("output")

    # GIVEN locations of current or specific folder
    output_filename = Path(location) / Path(input_file.replace(".json", f".{output_format}"))

    # WHEN writing transcript in any supported format
    tscribe.write(input_file, format=output_format, save_as=output_filename)

    # THEN check output exists
    assert output_filename.is_file()
    
    os.remove(output_filename)
